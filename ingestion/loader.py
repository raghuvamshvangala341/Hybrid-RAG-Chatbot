from pathlib import Path
from typing import List

from langchain_core.documents import Document

from pypdf import PdfReader
from docx import Document as DocxDocument


class DocumentLoader:

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".txt",
    }

    def load(self, file_path: str | Path) -> List[Document]:

        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"{file_path} not found.")

        extension = file_path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")

        if extension == ".pdf":
            return self._load_pdf(file_path)

        if extension == ".docx":
            return self._load_docx(file_path)

        if extension == ".txt":
            return self._load_txt(file_path)

        return []

    def _load_pdf(self, file_path: Path) -> List[Document]:

        reader = PdfReader(file_path)

        documents = []

        for page_number, page in enumerate(reader.pages, start=1):

            text = page.extract_text()

            if not text:
                continue

            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "file_type": "pdf",
                        "page": page_number,
                    },
                )
            )

        return documents

    def _load_docx(self, file_path: Path) -> List[Document]:

        doc = DocxDocument(file_path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        )

        return [
            Document(
                page_content=text,
                metadata={
                    "source": file_path.name,
                    "file_type": "docx",
                },
            )
        ]

    def _load_txt(self, file_path: Path) -> List[Document]:

        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read()

        return [
            Document(
                page_content=text,
                metadata={
                    "source": file_path.name,
                    "file_type": "txt",
                },
            )
        ]